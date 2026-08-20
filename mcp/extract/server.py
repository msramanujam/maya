"""MCP server: extract text from PDF, DOCX, XLSX and plain files.

Why this exists rather than an off-the-shelf server: the LibreChat image
is Alpine/musl, and every packaged extractor we tried pulls native wheels
built for glibc (`onnxruntime` has no musl build), so none of them
resolve here. pypdf, python-docx and openpyxl are pure Python and install
in under a second.

Access is bounded to the same roots the filesystem server uses. That is
belt and braces — /documents is read-only at the mount — but this server
could otherwise read anything in the container, including config.

Run: uvx --with pypdf --with python-docx --with openpyxl python server.py
"""
import json
import os
import sys

ALLOWED_ROOTS = ("/documents", "/projects")
MAX_CHARS = 200_000          # a whole book would otherwise flood the context


def _resolve(path):
    """Absolute, symlink-free path — or raise if it escapes the roots."""
    real = os.path.realpath(path)
    for root in ALLOWED_ROOTS:
        if real == root or real.startswith(root + os.sep):
            return real
    raise PermissionError(
        "path outside allowed directories: %s not in %s"
        % (path, ", ".join(ALLOWED_ROOTS))
    )


def _pdf(path):
    from pypdf import PdfReader
    return "\n".join((page.extract_text() or "") for page in PdfReader(path).pages)


def _docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append("# sheet: %s" % ws.title)
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                out.append("\t".join("" if c is None else str(c) for c in row))
    return "\n".join(out)


def _plain(path):
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read()


EXTRACTORS = {".pdf": _pdf, ".docx": _docx, ".xlsx": _xlsx}


def extract(path):
    real = _resolve(path)
    if not os.path.isfile(real):
        raise FileNotFoundError("not a file: %s" % path)
    ext = os.path.splitext(real)[1].lower()
    text = EXTRACTORS.get(ext, _plain)(real)
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[truncated at %d characters]" % MAX_CHARS
    return text


TOOLS = [{
    "name": "extract_document",
    "description": (
        "Extract readable text from a document. Handles PDF, DOCX and XLSX, "
        "and falls back to reading plain text. Only paths under /documents "
        "or /projects are accessible."
    ),
    "inputSchema": {
        "type": "object",
        "properties": {"path": {"type": "string",
                                "description": "Absolute path, e.g. /documents/report.pdf"}},
        "required": ["path"],
    },
}]


def reply(msg_id, result=None, error=None):
    out = {"jsonrpc": "2.0", "id": msg_id}
    if error is not None:
        out["error"] = error
    else:
        out["result"] = result
    sys.stdout.write(json.dumps(out) + "\n")
    sys.stdout.flush()


def main():
    for line in sys.stdin:
        line = line.strip()
        if not line:
            continue
        try:
            req = json.loads(line)
        except ValueError:
            continue
        method, msg_id = req.get("method"), req.get("id")

        if method == "initialize":
            reply(msg_id, {
                "protocolVersion": "2024-11-05",
                "capabilities": {"tools": {"listChanged": False}},
                "serverInfo": {"name": "maya-extract", "version": "1.0.0"},
            })
        elif method == "tools/list":
            reply(msg_id, {"tools": TOOLS})
        elif method == "tools/call":
            params = req.get("params") or {}
            if params.get("name") != "extract_document":
                reply(msg_id, error={"code": -32601,
                                     "message": "unknown tool: %s" % params.get("name")})
                continue
            try:
                text = extract((params.get("arguments") or {}).get("path", ""))
                reply(msg_id, {"content": [{"type": "text", "text": text}]})
            except Exception as exc:            # reported to the model, not swallowed
                reply(msg_id, {"content": [{"type": "text",
                                            "text": "%s: %s" % (type(exc).__name__, exc)}],
                               "isError": True})
        elif msg_id is not None:
            reply(msg_id, error={"code": -32601, "message": "unknown method: %s" % method})


if __name__ == "__main__":
    main()
